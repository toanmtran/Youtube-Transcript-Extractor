import os
import docx
import googleapiclient.discovery
from abc import ABC, abstractmethod
from google.oauth2.service_account import Credentials
from src.config import GOOGLE_DOCS_CREDENTIALS_PATH


class DocWriter(ABC):
    """Abstract base class for document writers."""

    @abstractmethod
    def write_video(self, title, content):
        """Writes the content of a single video to the document."""
        pass

    @abstractmethod
    def save(self):
        """Finalizes and saves the document."""
        pass


class GoogleDocsWriter(DocWriter):
    """Writes transcription data to a Google Doc"""

    def __init__(self, doc_link):
        self.doc_id = self._get_doc_id(doc_link)
        self.doc_service = self._authenticate()
        if not self.doc_id:
            raise ValueError("Invalid Google Doc link. Could not extract Document ID.")
        if not self.doc_service:
            raise ValueError("Failed to authenticate with Google Docs API.")

        # Test the connection and permissions
        try:
            self.doc_service.documents().get(documentId=self.doc_id).execute()
            print("Successfully connected to Google Docs.")
        except Exception as e:
            raise ValueError(f"Cannot access Google Doc. Check link and permissions. Error: {e}")

    def _get_doc_id(self, link):
        try:
            # The ID is between /d/ and /edit
            return link.split('/d/')[1].split('/')[0]
        except IndexError:
            return None

    def _authenticate(self):
        try:
            scope = ['https://www.googleapis.com/auth/documents', 'https://www.googleapis.com/auth/drive']
            creds = Credentials.from_service_account_file(GOOGLE_DOCS_CREDENTIALS_PATH, scopes=scope)
            return googleapiclient.discovery.build('docs', 'v1', credentials=creds)
        except Exception as e:
            print(f"Error authenticating with Google Docs API: {e}")
            return None

    def _text_length(self, text):
        # Google Docs API counts in UTF-16 code units.
        return len(text.encode('utf-16le')) // 2

    def write_video(self, title, content):
        # Always insert at index 1 to add new content to the top.
        insertion_point = 1

        # Prepare the full text to be inserted in one go.
        full_text = f"{title}\n{content}\n\n"

        # Calculate lengths for formatting ranges
        title_len = self._text_length(title)
        content_len = self._text_length(content)

        requests = [
            # 1. Insert all the text (title + content) at once.
            {
                'insertText': {
                    'location': {'index': insertion_point},
                    'text': full_text
                }
            },
            # 2. Format the title.
            {
                'updateParagraphStyle': {
                    'range': {
                        'startIndex': insertion_point,
                        'endIndex': insertion_point + title_len
                    },
                    'paragraphStyle': {'namedStyleType': 'HEADING_1'},
                    'fields': 'namedStyleType'
                }
            },
            {
                'updateTextStyle': {
                    'range': {
                        'startIndex': insertion_point,
                        'endIndex': insertion_point + title_len
                    },
                    'textStyle': {
                        'bold': True,
                        'fontSize': {'magnitude': 14, 'unit': 'PT'}
                    },
                    'fields': 'bold, fontSize'
                }
            },
            # 3. Format the body content.
            {
                'updateParagraphStyle': {
                    'range': {
                        'startIndex': insertion_point + title_len + 1,  # +1 for the newline after title
                        'endIndex': insertion_point + title_len + 1 + content_len
                    },
                    'paragraphStyle': {'namedStyleType': 'NORMAL_TEXT'},
                    'fields': 'namedStyleType'
                }
            },
            {
                'updateTextStyle': {
                    'range': {
                        'startIndex': insertion_point + title_len + 1,
                        'endIndex': insertion_point + title_len + 1 + content_len
                    },
                    'textStyle': {
                        'bold': False,
                        'fontSize': {'magnitude': 11, 'unit': 'PT'}
                    },
                    'fields': 'bold, fontSize'
                }
            }
        ]

        try:
            self.doc_service.documents().batchUpdate(
                documentId=self.doc_id, body={'requests': requests}
            ).execute()
        except Exception as e:
            if 'INVALID_ARGUMENT' in str(e) and 'exceeds the maximum' in str(e):
                print(f"Warning: Content for '{title}' is too long for a single Google Docs request and was skipped.")
            else:
                print(f"Error writing to Google Doc for video '{title}': {e}")

    def save(self):
        print("Google Doc has been updated.")


class MSWordWriter(DocWriter):
    """Writes transcription data to a Microsoft Word (.docx) file."""

    def __init__(self, storage_path):
        self.storage_path = storage_path
        self.doc = docx.Document()
        print("MS Word document initialized.")

    def write_video(self, title, content):
        """Appends a formatted title and content to the Word document."""
        try:
            # Filter out invalid XML characters that can crash python-docx
            clean_content = "".join(c for c in content if c.isprintable() or c in ('\n', '\t', '\r'))
            self.doc.add_heading(title, level=1)
            self.doc.add_paragraph(clean_content)
            self.doc.add_paragraph()  # Add a little space between entries
        except Exception as e:
            print(f"Error adding content for video '{title}' to Word doc: {e}")

    def save(self):
        """Saves the document to the specified path, avoiding overwrites."""
        base_filename = "YT_Captions.docx"
        doc_filename = base_filename
        count = 0

        while os.path.exists(os.path.join(self.storage_path, doc_filename)):
            count += 1
            doc_filename = f"YT_Captions ({count}).docx"

        full_path = os.path.join(self.storage_path, doc_filename)
        try:
            self.doc.save(full_path)
            print(f"\n📁 Saved results to {full_path}")
        except Exception as e:
            print(f"Error saving Word document: {e}")